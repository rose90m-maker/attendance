"""AI 안전진단 Blueprint — 현장 사진 → Ollama 비전 → 위험요인 추출 → DB 저장"""
import os, json, base64, re, math, urllib.request
from datetime import date, datetime
from functools import wraps
from flask import Blueprint, render_template, request, redirect, session, jsonify, send_file

safety_ai_bp = Blueprint("safety_ai", __name__, url_prefix="/safety_ai")

AI_HOST = os.environ.get("OLLAMA_HOST", "http://192.168.100.6:11434")
# 현장사진 8장 비교 결과 qwen 채택 (2026-07-24). gemma는 사람·보호구 미착용과
# 고소작업 추락 위험을 반복적으로 놓쳤고, 위험이 여럿인 현장에서도 항상 1건만 지적함.
# 되돌리려면 .env 에 AI_VISION_MODEL=gemma4:26b 지정.
VISION_MODEL = os.environ.get("AI_VISION_MODEL", "qwen2.5vl:32b")

RULE_KB = {
    "기계안전": "산업안전보건기준에 관한 규칙(방호장치·끼임 방지)",
    "전기안전": "산업안전보건기준에 관한 규칙(감전·충전부 방호)",
    "화재/소방": "산업안전보건기준에 관한 규칙(소화설비·인화성물질)",
    "통로/정리정돈": "산업안전보건기준에 관한 규칙(통로·정리정돈)",
    "보호구": "산업안전보건기준에 관한 규칙(개인보호구 지급·착용)",
    "추락/전도": "산업안전보건기준에 관한 규칙(추락 방지·안전난간)",
    "화학물질": "산업안전보건기준에 관한 규칙(관리대상 유해물질·MSDS)",
    "작업환경": "산업안전보건기준에 관한 규칙(환기·분진·소음·조도)",
}

SAFETY_VISION_SYSTEM = (
    "당신은 제조업 사업장의 산업안전 점검 전문가입니다. 업로드된 현장 사진을 분석해 위험요인을 추출하십시오.\n"
    "【절대 규칙 — 환각 금지】\n"
    "1) 사진에 실제로 보이는 물체·상황만 서술하라. 보이지 않는 것을 추정·상상해 쓰면 안 된다.\n"
    "2) 물체를 언급하기 전에 스스로 검증하라: \"이 물체가 사진의 어느 위치에 있는가?\" 답할 수 없으면 쓰지 마라.\n"
    "3) 사람이 보이지 않으면 보호구 착용 여부를 언급하지 마라. 사람이 명확히 보일 때만 "
    "안전모·안전화·보안경·방진마스크·귀마개·장갑·안전대 착용을 확인하고, 미착용 시 '보호구' 위험요인으로 등록.\n"
    "4) detected_objects에는 사진에 확실히 보이는 물체만 나열하라.\n"
    "5) 확신 없는 항목은 아예 제외하라. 적은 수의 확실한 지적이 많은 수의 추측보다 낫다.\n"
    "점검 관점(보이는 범위에서만): 기계 회전체·끼임부 방호덮개 / 전선·충전부·분전반 / 소화기·비상구·인화물 / "
    "통로·바닥(전선·기름·적치물) / 추락·개구부 / 화학물질·MSDS / 분진·환기·조도.\n"
    "위험등급 산정: 중대성 상+가능성 중이상=상, 중대성 중+가능성 중=중, 중대성 하+가능성 하=하. "
    "사망·끼임·추락·감전·화재 위험은 기본 '중' 이상으로 평가.\n"
    "주의: 사진에서 명확히 보이는 것과 추정되는 것을 구분(visibility=명확/추정, 추정이면 site_state에 '추정' 명시). "
    "작업자 얼굴·회사명 등 개인정보는 익명 처리. "
    "법규 조항 번호는 단정하지 말 것. 결과는 반드시 아래 JSON만 출력(설명·마크다운 금지, 한국어):\n"
    "{\n"
    '  "overall": "안전|주의|위험",\n'
    '  "summary": "종합 의견 2~3문장",\n'
    '  "image_summary": "사진 한 문장 설명",\n'
    '  "workplace_type": "작업장 유형",\n'
    '  "detected_objects": ["사다리","전선","소화기"],\n'
    '  "hazards": [\n'
    "    {\n"
    '      "name": "위험요인명",\n'
    '      "category": "기계안전|전기안전|화재/소방|통로/정리정돈|보호구|추락/전도|화학물질|작업환경",\n'
    '      "visibility": "명확|추정",\n'
    '      "likelihood": "상|중|하",  "severity": "상|중|하",  "grade": "상|중|하",\n'
    '      "accident_type": "끼임/감전/화재 등",\n'
    '      "site_state": "현장 상태 서술",\n'
    '      "immediate_action": "즉시조치",\n'
    '      "short_term_action": "단기조치",\n'
    '      "prevention_action": "재발방지조치",\n'
    '      "report_sentence": "보고서용 문장"\n'
    "    }\n"
    "  ]\n"
    "}\n"
    "위험요인이 없으면 hazards를 빈 배열로, overall은 '안전'으로."
)

_GRADE_ORD = {"상": 0, "중": 1, "하": 2}


def _conn():
    from app_maria import MARIA
    import pymysql
    return pymysql.connect(**MARIA)


def _compress_image_for_db(raw: bytes, max_kb: int = 600) -> bytes:
    """PIL이 있으면 ~600KB 이하로 압축(b64 후에도 max_allowed_packet 1MB 안), 없으면 원본 반환."""
    try:
        from PIL import Image
        import io as _io
        target = max_kb * 1024
        if len(raw) <= target:
            return raw
        img = Image.open(_io.BytesIO(raw))
        if img.width > 1920 or img.height > 1920:
            img.thumbnail((1920, 1920), Image.LANCZOS)
        for q in [80, 65, 50, 35, 25]:
            buf = _io.BytesIO()
            img.convert("RGB").save(buf, format="JPEG", quality=q)
            if buf.tell() <= target:
                return buf.getvalue()
        buf = _io.BytesIO()
        img.convert("RGB").save(buf, format="JPEG", quality=20)
        return buf.getvalue()
    except Exception:
        return raw


def _resize_for_vision(raw: bytes, max_px: int = 1280) -> bytes:
    """비전 모델 입력용 축소. 용량이 아니라 '해상도'를 제한한다.

    원본을 그대로 넘기면 고해상도 사진(아이폰 등)에서 비전 토큰이 폭증해
    GPU 메모리 부족으로 진단이 실패한다 (qwen2.5vl:32b, cudaMalloc OOM 확인).
    1280px 이하로 줄이면 정상 동작하며 위험요인 인식률 저하도 없었다.
    """
    try:
        from PIL import Image
        import io as _io
        img = Image.open(_io.BytesIO(raw))
        if img.width <= max_px and img.height <= max_px:
            return raw
        img.thumbnail((max_px, max_px), Image.LANCZOS)
        buf = _io.BytesIO()
        img.convert("RGB").save(buf, format="JPEG", quality=85)
        return buf.getvalue()
    except Exception:
        return raw


def _login_required(f):
    @wraps(f)
    def deco(*a, **kw):
        if not session.get("user_id"):
            return redirect("/login")
        return f(*a, **kw)
    return deco


def init_safety_ai_db(app):
    with app.app_context():
        conn = _conn(); cur = conn.cursor()
        cur.execute("""
            CREATE TABLE IF NOT EXISTS `safety_ai_inspections` (
                `id` INT AUTO_INCREMENT PRIMARY KEY,
                `site_name` VARCHAR(200) DEFAULT '',
                `inspector` VARCHAR(100) DEFAULT '',
                `dept` VARCHAR(100) DEFAULT '',
                `note` TEXT,
                `inspect_date` DATE,
                `image_mime` VARCHAR(50) DEFAULT 'image/jpeg',
                `image_b64` LONGTEXT,
                `overall` VARCHAR(20) DEFAULT '',
                `summary` TEXT,
                `image_summary` TEXT,
                `workplace_type` VARCHAR(100) DEFAULT '',
                `detected_json` TEXT,
                `hazards_json` LONGTEXT,
                `model` VARCHAR(100) DEFAULT '',
                `confirmed` TINYINT(1) DEFAULT 0,
                `created_at` DATETIME DEFAULT CURRENT_TIMESTAMP
            ) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4
        """)
        conn.commit(); conn.close()


def _vision_call(b64: str) -> dict:
    payload = {
        "model": VISION_MODEL,
        "messages": [
            {"role": "system", "content": SAFETY_VISION_SYSTEM},
            {"role": "user", "content": "이 현장 사진을 안전진단 해줘.", "images": [b64]},
        ],
        "stream": False,
        "format": "json",
        # 현장사진 실측으로 결정한 값 (2026-07-24). 같은 사진 5회 반복 → 결과 100% 동일.
        #  temperature 0 / top_k 1 / seed 고정 : 같은 사진엔 같은 진단이 나오게. 0.1 에서는
        #    재실행 시 '고소작업 추락 위험' 같은 핵심 항목이 통째로 사라지는 편차가 있었다.
        #  num_ctx 8192 : 기본 32768 이면 KV 캐시가 VRAM 을 잠식해 모델 27.1GB 중 4.4GB 가
        #    CPU 로 밀려나 2배 느려진다. 8192 면 전부 GPU 에 올라가 43초 → 20초.
        #    실측 최대 사용량은 2,702 토큰(입력 2,037 + 출력 665)이라 8192 는 3배 여유.
        "options": {
            "temperature": 0,
            "top_k": 1,
            "top_p": 1.0,
            "seed": 42,
            "num_ctx": 8192,
        },
        "think": False,
    }
    body = json.dumps(payload).encode("utf-8")
    req = urllib.request.Request(AI_HOST + "/api/chat", data=body,
                                 headers={"Content-Type": "application/json"})
    with urllib.request.urlopen(req, timeout=300) as r:
        d = json.loads(r.read())
    raw = ((d.get("message") or {}).get("content") or "").strip()
    t = re.sub(r"^```(?:json)?|```$", "", raw, flags=re.M).strip()
    try:
        return json.loads(t)
    except Exception:
        m = re.search(r"\{.*\}", t, re.S)
        if m:
            return json.loads(m.group(0))
        raise


def _enrich_hazards(hazards: list) -> list:
    for h in hazards:
        cat = h.get("category", "")
        rule = RULE_KB.get(cat)
        h["refs"] = {"rule": rule} if rule else {}
    hazards.sort(key=lambda h: _GRADE_ORD.get(h.get("grade", "하"), 3))
    return hazards


# ── 목록 + 업로드 ─────────────────────────────────────────────────────
@safety_ai_bp.route("", methods=["GET", "POST"])
@_login_required
def index():
    error = None
    if request.method == "POST":
        f = request.files.get("photo")
        site = (request.form.get("site_name") or "").strip()
        inspector = (request.form.get("inspector") or "").strip()
        idate = (request.form.get("inspect_date") or "").strip()
        dept = (request.form.get("dept") or "").strip()
        note = (request.form.get("note") or "").strip()
        if not f or not f.filename:
            error = "사진을 선택해 주세요."
        else:
            raw = f.read()
            if len(raw) > 12 * 1024 * 1024:
                error = "이미지가 너무 큽니다 (12MB 이하)."
            else:
                # AI 입력은 축소본을 쓴다 (원본 그대로 보내면 고해상도 사진에서 GPU OOM)
                b64 = base64.b64encode(_resize_for_vision(raw)).decode()
                try:
                    data = _vision_call(b64)
                    hazards = _enrich_hazards(data.get("hazards") or [])
                    ins_date = date.fromisoformat(idate) if idate else date.today()
                    b64_store = base64.b64encode(_compress_image_for_db(raw)).decode()
                    conn = _conn(); cur = conn.cursor()
                    cur.execute("""
                        INSERT INTO safety_ai_inspections
                        (site_name, inspector, dept, note, inspect_date, image_mime, image_b64,
                         overall, summary, image_summary, workplace_type, detected_json, hazards_json, model)
                        VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)
                    """, (
                        site or "(미입력)", inspector or "(미입력)", dept, note,
                        ins_date, f.mimetype or "image/jpeg", b64_store,
                        data.get("overall") or "검토필요",
                        data.get("summary") or "",
                        data.get("image_summary") or "",
                        data.get("workplace_type") or "",
                        json.dumps(data.get("detected_objects") or [], ensure_ascii=False),
                        json.dumps(hazards, ensure_ascii=False),
                        VISION_MODEL,
                    ))
                    conn.commit()
                    new_id = cur.lastrowid
                    conn.close()
                    return redirect(f"/safety_ai/{new_id}")
                except Exception as e:
                    error = f"AI 진단 실패: {e}"

    conn = _conn(); cur = conn.cursor()
    cur.execute("""SELECT id, site_name, inspector, dept, inspect_date, overall,
                          hazards_json, confirmed, created_at
                   FROM safety_ai_inspections ORDER BY id DESC LIMIT 30""")
    rows = cur.fetchall()
    conn.close()
    recent = []
    for r in rows:
        haz_cnt = len(json.loads(r[6] or "[]"))
        recent.append({
            "id": r[0], "site_name": r[1], "inspector": r[2], "dept": r[3],
            "inspect_date": r[4], "overall": r[5], "hazard_count": haz_cnt,
            "confirmed": r[7], "created_at": r[8],
        })
    return render_template("safety_ai_inspect.html", error=error, recent=recent,
                           today=date.today().isoformat(), vision_model=VISION_MODEL)


# ── 통계 (반드시 <int:sid> 라우트 전에 등록) ───────────────────────
@safety_ai_bp.route("/stats")
@_login_required
def stats():
    conn = _conn(); cur = conn.cursor()
    cur.execute("SELECT overall, hazards_json, inspect_date, site_name FROM safety_ai_inspections ORDER BY inspect_date")
    rows = cur.fetchall()
    conn.close()

    by_overall = {"안전": 0, "주의": 0, "위험": 0}
    by_grade = {"상": 0, "중": 0, "하": 0}
    by_cat = {}
    by_month = {}
    by_site = {}
    total_haz = 0

    for overall, hz_json, ins_date, site_name in rows:
        by_overall[overall] = by_overall.get(overall, 0) + 1
        ym = ins_date.strftime("%Y-%m") if ins_date else "-"
        m = by_month.setdefault(ym, {"recs": 0, "haz": 0, "high": 0})
        m["recs"] += 1
        site = site_name or "(미입력)"
        s = by_site.setdefault(site, {"recs": 0, "haz": 0, "high": 0, "last": None})
        s["recs"] += 1
        if ins_date and (s["last"] is None or ins_date > s["last"]):
            s["last"] = ins_date
        try:
            hz = json.loads(hz_json or "[]")
        except Exception:
            hz = []
        for h in hz:
            total_haz += 1
            g = h.get("grade", "하")
            by_grade[g] = by_grade.get(g, 0) + 1
            cat = h.get("category", "기타")
            by_cat[cat] = by_cat.get(cat, 0) + 1
            m["haz"] += 1; s["haz"] += 1
            if g == "상":
                m["high"] += 1; s["high"] += 1

    months = sorted(by_month.keys())
    site_rows = sorted(({"site": k, **v} for k, v in by_site.items()), key=lambda x: -x["high"])
    cat_rows = sorted(({"cat": k, "n": v} for k, v in by_cat.items()), key=lambda x: -x["n"])
    return render_template("safety_ai_stats.html",
                           total_recs=len(rows), total_haz=total_haz,
                           by_overall=by_overall, by_grade=by_grade,
                           months=months, by_month=by_month,
                           site_rows=site_rows, cat_rows=cat_rows)


# ── 보고서 ────────────────────────────────────────────────────────────
@safety_ai_bp.route("/<int:sid>")
@_login_required
def report(sid):
    conn = _conn(); cur = conn.cursor()
    cur.execute("SELECT * FROM safety_ai_inspections WHERE id=%s", (sid,))
    row = cur.fetchone()
    conn.close()
    if not row:
        return "진단을 찾을 수 없습니다.", 404
    cols = ["id","site_name","inspector","dept","note","inspect_date","image_mime","image_b64",
            "overall","summary","image_summary","workplace_type","detected_json","hazards_json",
            "model","confirmed","created_at"]
    rec = dict(zip(cols, row))
    hazards = json.loads(rec["hazards_json"] or "[]")
    detected = json.loads(rec["detected_json"] or "[]")
    return render_template("safety_ai_report.html", rec=rec, hazards=hazards, detected=detected)


# ── 수정 + 확정 ──────────────────────────────────────────────────────
@safety_ai_bp.route("/<int:sid>/edit", methods=["GET", "POST"])
@_login_required
def edit(sid):
    conn = _conn(); cur = conn.cursor()
    cur.execute("SELECT * FROM safety_ai_inspections WHERE id=%s", (sid,))
    row = cur.fetchone()
    if not row:
        conn.close()
        return "진단을 찾을 수 없습니다.", 404
    cols = ["id","site_name","inspector","dept","note","inspect_date","image_mime","image_b64",
            "overall","summary","image_summary","workplace_type","detected_json","hazards_json",
            "model","confirmed","created_at"]
    rec = dict(zip(cols, row))
    hazards = json.loads(rec["hazards_json"] or "[]")

    if request.method == "POST":
        new_overall = (request.form.get("overall") or rec["overall"]).strip()
        new_summary = (request.form.get("summary") or "").strip()
        for i, h in enumerate(hazards):
            h["name"] = (request.form.get(f"name_{i}") or h.get("name", "")).strip()
            h["grade"] = (request.form.get(f"grade_{i}") or h.get("grade", "")).strip()
            h["accident_type"] = (request.form.get(f"acc_{i}") or h.get("accident_type", "")).strip()
            h["site_state"] = (request.form.get(f"site_{i}") or "").strip()
            h["immediate_action"] = (request.form.get(f"imm_{i}") or "").strip()
            h["short_term_action"] = (request.form.get(f"short_{i}") or "").strip()
            h["prevention_action"] = (request.form.get(f"prev_{i}") or "").strip()
            h["report_sentence"] = (request.form.get(f"rep_{i}") or "").strip()
        dels = set(request.form.getlist("delete"))
        hazards = [h for i, h in enumerate(hazards) if str(i) not in dels]
        hazards.sort(key=lambda h: _GRADE_ORD.get(h.get("grade", "하"), 3))
        confirmed = 1 if request.form.get("action") == "confirm" else rec["confirmed"]
        cur.execute("""UPDATE safety_ai_inspections
                       SET overall=%s, summary=%s, hazards_json=%s, confirmed=%s
                       WHERE id=%s""",
                    (new_overall, new_summary, json.dumps(hazards, ensure_ascii=False),
                     confirmed, sid))
        conn.commit(); conn.close()
        return redirect(f"/safety_ai/{sid}")

    conn.close()
    return render_template("safety_ai_edit.html", rec=rec, hazards=hazards)


# ── DOCX 다운로드 ────────────────────────────────────────────────────
@safety_ai_bp.route("/<int:sid>/docx")
@_login_required
def docx(sid):
    import io
    try:
        from docx import Document
        from docx.shared import Inches
    except ImportError:
        return "python-docx 미설치", 500
    conn = _conn(); cur = conn.cursor()
    cur.execute("SELECT * FROM safety_ai_inspections WHERE id=%s", (sid,))
    row = cur.fetchone()
    conn.close()
    if not row:
        return "진단을 찾을 수 없습니다.", 404
    cols = ["id","site_name","inspector","dept","note","inspect_date","image_mime","image_b64",
            "overall","summary","image_summary","workplace_type","detected_json","hazards_json",
            "model","confirmed","created_at"]
    rec = dict(zip(cols, row))
    hazards = json.loads(rec["hazards_json"] or "[]")

    doc = Document()
    doc.add_heading("사진 기반 안전진단 보고서", 0)
    doc.add_paragraph(f"현장명: {rec['site_name']}    점검일: {rec['inspect_date']}    점검자: {rec['inspector']}")
    doc.add_paragraph(f"AI 모델: {rec['model']}    종합판정: {rec['overall']}")
    doc.add_paragraph("※ 본 결과는 확정 판정이 아닌 점검 보조 의견입니다.")
    try:
        img = io.BytesIO(base64.b64decode(rec["image_b64"]))
        doc.add_picture(img, width=Inches(4.5))
    except Exception:
        pass
    doc.add_heading("종합 의견", level=1)
    doc.add_paragraph(rec["summary"] or "")
    doc.add_heading("위험요인별 상세 진단", level=1)
    for i, h in enumerate(hazards, 1):
        doc.add_heading(f"No.{i} {h.get('name','')} [위험등급 {h.get('grade','')}]", level=2)
        doc.add_paragraph(f"유형: {h.get('category','')} / 사고유형: {h.get('accident_type','')} "
                          f"(가능성 {h.get('likelihood','')}, 중대성 {h.get('severity','')})")
        doc.add_paragraph(f"현장상태: {h.get('site_state','')}")
        refs = h.get("refs") or {}
        rule = refs.get("rule")
        doc.add_paragraph(f"관련기준: {rule or '관련 기준 검토 필요'}")
        doc.add_paragraph(f"즉시조치: {h.get('immediate_action','')}")
        if h.get("short_term_action"):
            doc.add_paragraph(f"단기조치: {h.get('short_term_action')}")
        if h.get("prevention_action"):
            doc.add_paragraph(f"재발방지: {h.get('prevention_action')}")
    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    fn = f"safety_{sid}_{rec['inspect_date']}.docx"
    return send_file(buf, as_attachment=True, download_name=fn,
                     mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
